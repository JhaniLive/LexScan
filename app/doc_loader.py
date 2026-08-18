import os
import re

# Phone cameras are the most common way a document arrives here, so the list
# covers what phones actually produce — iPhones save .heic by default, Windows
# and scanners produce .tif, and browsers hand over .jfif and .avif.
IMAGE_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp",
    ".heic", ".heif", ".tif", ".tiff", ".jfif", ".avif", ".jpe",
}

# Formats OpenCV (and so the OCR engine) cannot open directly — converted first.
NEEDS_CONVERSION = {".heic", ".heif", ".avif", ".jfif", ".gif", ".tif", ".tiff"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
TEXT_EXTENSIONS = {".txt", ".md", ".rtf", ".csv"}

SUPPORTED_EXTENSIONS = (
    IMAGE_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | TEXT_EXTENSIONS
)

# Anything under this goes straight to the agents; anything over gets condensed first.
DIRECT_ANALYSIS_LIMIT = 45_000
CHUNK_SIZE = 12_000
CHUNK_OVERLAP = 800

# Render resolution for OCR. 200 is the sweet spot — 150 starts losing small
# print, 300 doubles the time for no gain on typed documents.
OCR_DPI = 200

# A PDF with less text than this in its layer is treated as a scan.
TEXT_LAYER_FLOOR = 200


class UnsupportedDocument(Exception):
    pass


def extract_text_from_pdf(pdf_path):
    """Extract a PDF's text layer, page by page, tagging pages for citation.

    Returns "" for a scanned PDF — there is no text layer to find, and the
    caller falls back to OCR.
    """
    import pymupdf

    pages = []
    with pymupdf.open(pdf_path) as document:
        for number, page in enumerate(document, 1):
            page_text = page.get_text().strip()
            if page_text:
                pages.append(f"[Page {number}]\n{page_text}")
    return "\n\n".join(pages)


# The OCR engine loads its models on first use — keep the instance around.
_ocr_engine = None


def _get_ocr_engine():
    """Lazily build the local OCR engine. Returns None if it isn't installed."""
    global _ocr_engine
    if _ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
        except ImportError:
            return None
        _ocr_engine = RapidOCR()
    return _ocr_engine


def local_ocr_available():
    return _get_ocr_engine() is not None


def _ocr_bytes(engine, image_bytes):
    """Run OCR over one rendered page and join the recognised lines."""
    result, _ = engine(image_bytes)
    if not result:
        return ""
    return "\n".join(line[1] for line in result)


def ocr_pdf(pdf_path, dpi=OCR_DPI):
    """Render each page of a scanned PDF and read it with the local OCR engine."""
    import pymupdf

    engine = _get_ocr_engine()
    if engine is None:
        raise UnsupportedDocument("Local OCR is not installed.")

    pages = []
    with pymupdf.open(pdf_path) as document:
        for number, page in enumerate(document, 1):
            image = page.get_pixmap(dpi=dpi).tobytes("png")
            page_text = _ocr_bytes(engine, image).strip()
            if page_text:
                pages.append(f"[Page {number}]\n{page_text}")
    return "\n\n".join(pages)


def _image_bytes(image_path):
    """Return image bytes the OCR engine can read, converting formats it can't.

    iPhones save .heic, scanners produce .tif, browsers hand over .jfif — none of
    which OpenCV opens. Those get converted to PNG here. Very large phone photos
    are capped at 4000px, which costs nothing in accuracy and saves real time.
    """
    ext = os.path.splitext(image_path)[1].lower()

    if ext not in NEEDS_CONVERSION:
        with open(image_path, "rb") as f:
            data = f.read()
        # A format the engine reads natively only needs handling if it is huge.
        try:
            import io

            from PIL import Image

            with Image.open(io.BytesIO(data)) as image:
                if max(image.size) <= 4000:
                    return data
        except Exception:
            return data

    import io

    from PIL import Image

    if ext in {".heic", ".heif"}:
        try:
            import pillow_heif

            pillow_heif.register_heif_opener()
        except ImportError:
            raise UnsupportedDocument(
                "iPhone photos (.heic) need the `pillow-heif` package — install it "
                "with `pip install pillow-heif`, or share the photo as JPEG."
            )

    try:
        with Image.open(image_path) as image:
            image = image.convert("RGB")

            # Only ever shrink. Upscaling a small screenshot was measured to
            # make recognition worse, not better — interpolation smears glyphs
            # that are already degraded, and the engine reads the original
            # cleanly. Capping huge phone photos just saves time.
            longest = max(image.size)
            if longest > 4000:
                scale = 4000 / longest
                image = image.resize(
                    (int(image.width * scale), int(image.height * scale)),
                    Image.LANCZOS,
                )

            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            return buffer.getvalue()
    except UnsupportedDocument:
        raise
    except Exception as e:
        raise UnsupportedDocument(f"Could not open that image: {e}")


def ocr_image_file(image_path):
    """Read a photo or screenshot of a document with the local OCR engine."""
    engine = _get_ocr_engine()
    if engine is None:
        raise UnsupportedDocument("Local OCR is not installed.")

    return _ocr_bytes(engine, _image_bytes(image_path)).strip()


def extract_text_from_docx(docx_path):
    """Extract paragraphs and table cells from a .docx contract."""
    from docx import Document

    document = Document(docx_path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_plain(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def normalise(text):
    """Collapse the ragged whitespace PDF extraction tends to produce."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split long text on paragraph boundaries, with a little overlap between chunks."""
    if len(text) <= size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        if end >= len(text):
            chunks.append(text[start:])
            break

        # Prefer to break at a paragraph, then a sentence, then wherever we landed.
        window = text[start:end]
        split_at = window.rfind("\n\n")
        if split_at < size // 2:
            split_at = window.rfind(". ")
        if split_at < size // 2:
            split_at = size

        chunks.append(text[start:start + split_at])
        start += max(split_at - overlap, 1)

    return [c.strip() for c in chunks if c.strip()]


def describe_document(text):
    """Cheap stats shown to the user before the LLM ever sees the document."""
    words = len(text.split())
    return {
        "characters": len(text),
        "words": words,
        "pages": len(re.findall(r"\[Page \d+\]", text)) or max(1, words // 500),
        "needs_condensing": len(text) > DIRECT_ANALYSIS_LIMIT,
    }


NO_OCR_MESSAGE = (
    "This looks like a scan, and no OCR backend is available. Install the local "
    "one with `pip install pymupdf rapidocr-onnxruntime`, or set LLM_VISION_MODEL "
    "in .env to a model that can see."
)


async def load_document(path, name=None, ocr=None, on_status=None):
    """Read an uploaded file into plain text.

    Scans go through the local OCR engine when it is installed — it needs no
    model server and no vision model. `ocr` is an optional async fallback that
    reads images with a vision LLM.
    """
    import asyncio

    name = name or os.path.basename(path)
    ext = os.path.splitext(name)[1].lower()

    async def read_scan(reader, label):
        """OCR locally if we can, otherwise hand the file to the vision model."""
        if local_ocr_available():
            if on_status:
                await on_status(f"`{name}` is a scan — reading it with OCR...")
            return await asyncio.to_thread(reader, path), label
        if ocr:
            if on_status:
                await on_status(f"`{name}` is a scan — reading it with the vision model...")
            return await ocr(path), f"{label} via vision model"
        raise UnsupportedDocument(NO_OCR_MESSAGE)

    if ext in PDF_EXTENSIONS:
        text = await asyncio.to_thread(extract_text_from_pdf, path)
        source = "pdf"
        # A scanned contract has no text layer worth the name — OCR it instead.
        if len(text.strip()) < TEXT_LAYER_FLOOR:
            text, source = await read_scan(ocr_pdf, "scanned pdf (OCR)")

    elif ext in DOCX_EXTENSIONS:
        text = await asyncio.to_thread(extract_text_from_docx, path)
        source = "docx"

    elif ext in IMAGE_EXTENSIONS:
        text, source = await read_scan(ocr_image_file, "image (OCR)")

    elif ext in TEXT_EXTENSIONS or ext == "":
        text = extract_text_from_plain(path)
        source = "text"

    else:
        raise UnsupportedDocument(
            f"I can't read `{ext or name}`. What works: **PDF**, **DOCX**, "
            f"**TXT/MD**, and photos or scans as **JPG, PNG, HEIC, WEBP, TIFF, "
            f"BMP or GIF**. If it's a `.doc` (the old Word format) or `.pages`, "
            f"open it and save as PDF or DOCX first."
        )

    return normalise(text), source
